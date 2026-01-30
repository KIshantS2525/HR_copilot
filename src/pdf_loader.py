import os
from pypdf import PdfReader

try:
    from docx import Document
except ImportError:
    Document = None


def load_pdf(path: str) -> str:
    """Load text from a PDF file."""
    reader = PdfReader(path)
    text = ""

    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"

    return text


def load_docx(path: str) -> str:
    """Load text from a DOCX file."""
    if Document is None:
        raise ImportError("python-docx is required to load DOCX files. Install with `pip install python-docx`")
    
    doc = Document(path)
    text = ""
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    # also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    
    return text


def load_doc(path: str) -> str:
    """Load text from a DOC file using python-docx (limited support for older .doc files)."""
    # python-docx has limited support for older .doc format
    # For better .doc support, consider using python-docx2docx or other converters
    if Document is None:
        raise ImportError("python-docx is required to load DOC files. Install with `pip install python-docx`")
    
    try:
        doc = Document(path)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        return text
    except Exception as e:
        raise ValueError(f"Failed to load DOC file: {str(e)}")


def load_document(path: str) -> str:
    """Load text from PDF, DOCX, or DOC file based on extension."""
    _, ext = os.path.splitext(path)
    ext = ext.lower()
    
    if ext == ".pdf":
        return load_pdf(path)
    elif ext == ".docx":
        return load_docx(path)
    elif ext == ".doc":
        return load_doc(path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported formats: .pdf, .docx, .doc")
